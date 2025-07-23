"""
PDF and document parsing utilities for extracting text from resume files.
Supports PDF, DOCX, and TXT formats with multiple parsing strategies.
"""
import os
import logging
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
import mimetypes

# PDF parsing libraries
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# DOCX parsing
try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# Additional text extraction
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

from config.settings import settings

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors."""
    pass


class DocumentParser:
    """Main document parser class with multiple parsing strategies."""
    
    def __init__(self):
        """Initialize document parser with available libraries."""
        self.supported_formats = []
        
        # Check available libraries
        if HAS_PYPDF2 or HAS_PDFPLUMBER or HAS_PYMUPDF:
            self.supported_formats.append('.pdf')
        
        if HAS_PYTHON_DOCX:
            self.supported_formats.append('.docx')
            
        self.supported_formats.extend(['.txt', '.rtf'])
        
        logger.info(f"Document parser initialized. Supported formats: {self.supported_formats}")
        
        # Check which PDF libraries are available
        pdf_libs = []
        if HAS_PDFPLUMBER:
            pdf_libs.append("pdfplumber")
        if HAS_PYMUPDF:
            pdf_libs.append("PyMuPDF")
        if HAS_PYPDF2:
            pdf_libs.append("PyPDF2")
            
        logger.info(f"Available PDF libraries: {pdf_libs}")
    
    def extract_text(
        self, 
        file_path: Union[str, Path], 
        strategy: str = "auto"
    ) -> Dict[str, Any]:
        """
        Extract text from document with metadata.
        
        Args:
            file_path: Path to the document file
            strategy: Parsing strategy ('auto', 'pdfplumber', 'pymupdf', 'pypdf2')
        
        Returns:
            Dictionary with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentParsingError(f"File not found: {file_path}")
        
        # Validate file size
        file_size = file_path.stat().st_size
        if file_size > settings.max_file_size:
            raise DocumentParsingError(
                f"File too large: {file_size} bytes. Max allowed: {settings.max_file_size} bytes"
            )
        
        # Determine file type
        file_extension = file_path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Validate supported format
        if file_extension not in self.supported_formats:
            raise DocumentParsingError(
                f"Unsupported file format: {file_extension}. "
                f"Supported formats: {self.supported_formats}"
            )
        
        logger.info(f"Parsing document: {file_path.name} ({file_extension})")
        
        try:
            # Route to appropriate parser
            if file_extension == '.pdf':
                result = self._parse_pdf(file_path, strategy)
            elif file_extension == '.docx':
                result = self._parse_docx(file_path)
            elif file_extension in ['.txt', '.rtf']:
                result = self._parse_text(file_path)
            else:
                raise DocumentParsingError(f"No parser available for {file_extension}")
            
            # Add file metadata
            result.update({
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_size': file_size,
                'file_extension': file_extension,
                'mime_type': mime_type,
                'parsing_strategy': strategy
            })
            
            # Validate extracted text
            if not result.get('text') or len(result['text'].strip()) < 50:
                logger.warning(f"Very little text extracted from {file_path.name}")
                result['warnings'] = result.get('warnings', [])
                result['warnings'].append("Very little text was extracted from this document")
            
            logger.info(f"Successfully parsed {file_path.name}. "
                       f"Text length: {len(result.get('text', ''))} characters")
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            raise DocumentParsingError(f"Failed to parse document: {e}")
    
    def _parse_pdf(self, file_path: Path, strategy: str = "auto") -> Dict[str, Any]:
        """Parse PDF file using available libraries."""
        strategies = []
        
        if strategy == "auto":
            # Try strategies in order of preference
            if HAS_PDFPLUMBER:
                strategies.append("pdfplumber")
            if HAS_PYMUPDF:
                strategies.append("pymupdf")
            if HAS_PYPDF2:
                strategies.append("pypdf2")
        else:
            strategies = [strategy]
        
        last_error = None
        
        for strategy_name in strategies:
            try:
                logger.info(f"Trying PDF parsing with {strategy_name}")
                
                if strategy_name == "pdfplumber" and HAS_PDFPLUMBER:
                    return self._parse_pdf_pdfplumber(file_path)
                elif strategy_name == "pymupdf" and HAS_PYMUPDF:
                    return self._parse_pdf_pymupdf(file_path)
                elif strategy_name == "pypdf2" and HAS_PYPDF2:
                    return self._parse_pdf_pypdf2(file_path)
                else:
                    continue
                    
            except Exception as e:
                logger.warning(f"PDF parsing failed with {strategy_name}: {e}")
                last_error = e
                continue
        
        if last_error:
            raise DocumentParsingError(f"All PDF parsing strategies failed. Last error: {last_error}")
        else:
            raise DocumentParsingError("No PDF parsing libraries available")
    
    def _parse_pdf_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF using pdfplumber (best for tables and layout)."""
        import pdfplumber
        
        text_parts = []
        tables = []
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                'page_count': len(pdf.pages),
                'parsing_method': 'pdfplumber',
                'creator': pdf.metadata.get('Creator', ''),
                'producer': pdf.metadata.get('Producer', ''),
                'creation_date': pdf.metadata.get('CreationDate', '')
            }
            
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                
                # Extract tables
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        tables.append({
                            'page': page_num,
                            'data': table
                        })
        
        return {
            'text': '\n\n'.join(text_parts),
            'tables': tables,
            'metadata': metadata,
            'warnings': []
        }
    
    def _parse_pdf_pymupdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF using PyMuPDF (good for complex layouts)."""
        import fitz
        
        doc = fitz.open(str(file_path))
        text_parts = []
        
        metadata = {
            'page_count': doc.page_count,
            'parsing_method': 'pymupdf',
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'creator': doc.metadata.get('creator', '')
        }
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        return {
            'text': '\n\n'.join(text_parts),
            'tables': [],  # PyMuPDF doesn't extract tables as easily
            'metadata': metadata,
            'warnings': []
        }
    
    def _parse_pdf_pypdf2(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF using PyPDF2 (basic text extraction)."""
        import PyPDF2
        
        text_parts = []
        metadata = {}
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            metadata = {
                'page_count': len(pdf_reader.pages),
                'parsing_method': 'pypdf2'
            }
            
            # Try to get metadata
            try:
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    })
            except:
                pass
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
        
        return {
            'text': '\n\n'.join(text_parts),
            'tables': [],
            'metadata': metadata,
            'warnings': []
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file using python-docx."""
        if not HAS_PYTHON_DOCX:
            raise DocumentParsingError("python-docx library not available for DOCX parsing")
        
        from docx import Document
        
        doc = Document(str(file_path))
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Only add non-empty rows
                    table_data.append(row_data)
            if table_data:
                tables.append({'data': table_data})
        
        # Basic metadata
        metadata = {
            'paragraph_count': len(paragraphs),
            'table_count': len(tables),
            'parsing_method': 'python-docx'
        }
        
        # Try to get document properties
        try:
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            })
        except:
            pass
        
        return {
            'text': '\n\n'.join(paragraphs),
            'tables': tables,
            'metadata': metadata,
            'warnings': []
        }
    
    def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text files."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                
                return {
                    'text': text,
                    'tables': [],
                    'metadata': {
                        'encoding': encoding,
                        'parsing_method': 'text',
                        'line_count': len(text.splitlines())
                    },
                    'warnings': []
                }
                
            except UnicodeDecodeError:
                continue
        
        raise DocumentParsingError(f"Could not decode text file with any of: {encodings}")
    
    def extract_text_from_bytes(
        self, 
        file_bytes: bytes, 
        filename: str,
        strategy: str = "auto"
    ) -> Dict[str, Any]:
        """Extract text from file bytes (useful for web uploads)."""
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_file_path = temp_file.name
        
        try:
            result = self.extract_text(temp_file_path, strategy)
            result['original_filename'] = filename
            return result
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
            except:
                pass
    
    def validate_resume_content(self, text: str) -> Dict[str, Any]:
        """Validate that extracted text looks like a resume."""
        text_lower = text.lower()
        
        # Common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'bachelor', 'master',
            'email', 'phone', 'address', 'linkedin', 'github'
        ]
        
        found_indicators = [indicator for indicator in resume_indicators 
                          if indicator in text_lower]
        
        # Email pattern check
        import re
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        has_email = bool(re.search(email_pattern, text))
        
        # Phone pattern check
        phone_pattern = r'[\+]?[\d\s\-\(\)]{10,}'
        has_phone = bool(re.search(phone_pattern, text))
        
        confidence_score = len(found_indicators) / len(resume_indicators)
        
        validation_result = {
            'is_likely_resume': confidence_score > 0.2 or has_email,
            'confidence_score': confidence_score,
            'found_indicators': found_indicators,
            'has_email': has_email,
            'has_phone': has_phone,
            'text_length': len(text),
            'word_count': len(text.split())
        }
        
        return validation_result


# Global parser instance
document_parser = DocumentParser()


# Convenience functions
def extract_text_from_file(file_path: Union[str, Path], strategy: str = "auto") -> str:
    """Extract text from file and return just the text content."""
    result = document_parser.extract_text(file_path, strategy)
    return result['text']


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract text from bytes and return just the text content."""
    result = document_parser.extract_text_from_bytes(file_bytes, filename)
    return result['text']


def parse_resume_file(file_path: Union[str, Path]) -> Dict[str, Any]:
    """Parse resume file with validation."""
    result = document_parser.extract_text(file_path)
    
    # Add resume validation
    validation = document_parser.validate_resume_content(result['text'])
    result['validation'] = validation
    
    if not validation['is_likely_resume']:
        logger.warning(f"File {file_path} may not be a resume (confidence: {validation['confidence_score']:.2f})")
    
    return result


def get_supported_formats() -> List[str]:
    """Get list of supported file formats."""
    return document_parser.supported_formats.copy()